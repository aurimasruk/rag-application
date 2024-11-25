import os
import logging
from dotenv import load_dotenv
from pinecone import Pinecone, ServerlessSpec
from PyPDF2 import PdfReader
from fpdf import FPDF
import docx
import pandas as pd
import google.generativeai as genai
from llama_index.core import VectorStoreIndex, Document
from llama_index.llms.gemini import Gemini
from llama_index.embeddings.gemini import GeminiEmbedding
from llama_index.vector_stores.pinecone import PineconeVectorStore
from llama_index.core.ingestion import IngestionPipeline
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core.settings import Settings

# Setup logging
logging.basicConfig(level=logging.ERROR, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
GOOGLE_APPLICATION_CREDENTIALS = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")

# Validate environment variables
if not GOOGLE_APPLICATION_CREDENTIALS or not os.path.exists(GOOGLE_APPLICATION_CREDENTIALS):
    logger.error("Google Application Default Credentials not set or file not found.")
    raise ValueError("Please set the GOOGLE_APPLICATION_CREDENTIALS environment variable.")

if not GEMINI_API_KEY:
    logger.error("GEMINI_API_KEY not set.")
    raise ValueError("Please set the GEMINI_API_KEY environment variable.")

if not PINECONE_API_KEY:
    logger.error("PINECONE_API_KEY not set.")
    raise ValueError("Please set the PINECONE_API_KEY environment variable.")

# Configure Google Generative AI and LLM
genai.configure(api_key=GEMINI_API_KEY)
gemini_llm = Gemini(api_key=GEMINI_API_KEY)
embedding_model = GeminiEmbedding()

# Configure Pinecone
pinecone = Pinecone(api_key=PINECONE_API_KEY)
index_name = "demo"

# Check or create Pinecone index
existing_indexes = [idx["name"] for idx in pinecone.list_indexes()]
if index_name not in existing_indexes:
    logger.info(f"Creating new Pinecone index: {index_name}")
    pinecone.create_index(
        name=index_name,
        dimension=768,
        metric="cosine",
        spec=ServerlessSpec(cloud="aws", region="us-east-1"),
    )
else:
    logger.info(f"Pinecone index '{index_name}' already exists. Proceeding with existing index.")

# Setup Vector Store
vector_store = PineconeVectorStore(index_name=index_name, api_key=PINECONE_API_KEY, environment="us-east-1")

# Set LlamaIndex settings
Settings.llm = gemini_llm
Settings.embed_model = embedding_model

def read_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        return "".join(page.extract_text() for page in reader.pages)
    except Exception as e:
        logger.error(f"Error reading PDF file: {file_path}, Error: {str(e)}")
        return ""

def read_word(file_path):
    try:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        logger.error(f"Error reading Word file: {file_path}, Error: {str(e)}")
        return ""

def read_excel(file_path):
    try:
        df = pd.read_excel(file_path)
        logger.info(f"Excel file data:\n{df.to_string(index=False)}")
        
        # Convert each row into descriptive text
        content = []
        for _, row in df.iterrows():
            row_description = ", ".join([f"{col}: {row[col]}" for col in df.columns if pd.notna(row[col])])
            content.append(row_description)
        
        return "\n".join(content)

    except Exception as e:
        logger.error(f"Error reading Excel file: {file_path}, Error: {str(e)}")
        return ""

def read_documents(folder_path):
    documents = []
    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)
        if file_name.endswith(".pdf"):
            documents.append(read_pdf(file_path))
        elif file_name.endswith(".docx"):
            documents.append(read_word(file_path))
        elif file_name.endswith(".xlsx"):
            documents.append(read_excel(file_path))
    return [Document(text=doc) for doc in documents if doc.strip()]

def ingest_documents(folder_path):
    documents = read_documents(folder_path)
    logger.info(f"Documents to ingest: {[doc.text for doc in documents]}")  # Log document texts
    pipeline = IngestionPipeline(
        transformations=[SentenceSplitter(chunk_size=1024), embedding_model],
        vector_store=vector_store,
    )
    pipeline.run(documents=documents, show_progress=False)
    logger.info("Documents successfully ingested into Pinecone.")
    return VectorStoreIndex.from_vector_store(vector_store=vector_store)

def query_index(index, query, top_k=5):
    system_prompt = "You are an AI assistant. Provide answers strictly based on the provided context."
    system_prompt += "\nExamples: \n- Q: What features does Product X have? A: Product X has 10W, Green features."

    retriever = VectorIndexRetriever(index=index, similarity_top_k=top_k, similarity_threshold = 0.7)       # Check threshold=0.8
    query_engine = RetrieverQueryEngine(retriever=retriever)
    return query_engine.query(f"{system_prompt}\n{query}")

def generate_test_documents(folder_path):
    os.makedirs(folder_path, exist_ok=True)
    
    pdf_path = os.path.join(folder_path, "product_descriptions.pdf")
    word_path = os.path.join(folder_path, "product_specifications.docx")
    excel_path = os.path.join(folder_path, "product_features.xlsx")
    
    if not os.path.exists(pdf_path):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("helvetica", size=12)
        pdf.cell(200, 10, text="Product Descriptions", new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.multi_cell(0, 10, text="Product A: Durable, suitable for outdoor use.\nProduct B: 3-year warranty.")
        pdf.output(pdf_path)
        logger.info("Generated PDF document.")

    if not os.path.exists(word_path):
        doc = docx.Document()
        doc.add_heading("Product Specifications", level=1)
        doc.add_paragraph("Product A: Energy usage - 15W, Color - White.")
        doc.add_paragraph("Product B: Energy usage - 20W, Color - Black.")
        doc.save(word_path)
        logger.info("Generated Word document.")

    if not os.path.exists(excel_path):
        data = {"Product": ["A", "B"], "Features": ["15W, White", "20W, Black"], "Warranty (Years)": [5, 3]}
        df = pd.DataFrame(data)
        df.to_excel(excel_path, index=False)
        logger.info("Generated Excel document.")

def main():
    folder_path = "test_documents"
    generate_test_documents(folder_path)
    logger.info("Starting document ingestion...")
    index = ingest_documents(folder_path)
    query = input("Enter your product-related question: ")
    response = query_index(index, query)
    print(f"\nQuestion: {query}")
    print(f"Answer: {response}")

if __name__ == "__main__":
    main()
